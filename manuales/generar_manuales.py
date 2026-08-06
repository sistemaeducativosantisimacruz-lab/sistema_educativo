# -*- coding: utf-8 -*-
"""
Generador de Manuales de Usuario - Sistema Educativo Santísima Cruz
Genera 4 archivos .docx:
  - Manual_General_Directorio.docx
  - Manual_Administrador.docx
  - Manual_Docente.docx
  - Manual_Estudiante.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import copy

OUTPUT_DIR = r"c:\xampp\htdocs\sistema_educativo\manuales"

# ─────────────────────────────────────────────────────────
# HELPERS DE FORMATO
# ─────────────────────────────────────────────────────────

def setup_document(doc, title_color=(0x1F, 0x35, 0x64)):
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(1.2)
    sec.right_margin = Inches(1.0)
    sec.top_margin  = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    styles = doc.styles
    n = styles['Normal']
    n.font.name = 'Calibri'; n.font.size = Pt(11)

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'; h1.font.size = Pt(18); h1.font.bold = True
    h1.font.color.rgb = RGBColor(*title_color)
    h1.paragraph_format.space_before = Pt(24); h1.paragraph_format.space_after = Pt(6)

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'; h2.font.size = Pt(14); h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    h2.paragraph_format.space_before = Pt(18); h2.paragraph_format.space_after = Pt(4)

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'; h3.font.size = Pt(12); h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1F, 0x7A, 0x8C)
    h3.paragraph_format.space_before = Pt(12); h3.paragraph_format.space_after = Pt(3)

    if 'Tip' not in [s.name for s in styles]:
        tip = styles.add_style('Tip', WD_STYLE_TYPE.PARAGRAPH)
        tip.font.name = 'Calibri'; tip.font.size = Pt(10); tip.font.italic = True

def body(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    p.paragraph_format.space_after = Pt(6)
    return p

def bold_body(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    return p

def step(doc, number, text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.2)
    run_num = p.add_run(f"Paso {number}: ")
    run_num.font.bold = True
    run_num.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    p.add_run(text)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
    p.paragraph_format.space_after = Pt(3)
    return p

def screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF9E6')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(f"📸  ESPACIO PARA CAPTURA:  {description}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x50, 0x00)
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '4'); b.set(qn('w:color'), 'FFC107')
        pBdr.append(b)
    p._p.get_or_add_pPr().append(pBdr)
    return p

def tip_box(doc, text, color='E8F4FD'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.italic = True
    return p

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), '2F5496')
    pBdr.append(b)
    p._p.get_or_add_pPr().append(pBdr)

def section_title_banner(doc, text, bg='2F5496'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), bg)
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(f"  {text}  ")
    run.font.name = 'Calibri'; run.font.size = Pt(13); run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return p

def cover_page(doc, title, subtitle, role_color, version="1.0"):
    setup_document(doc, title_color=role_color)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    run = p.add_run("INSTITUCIÓN EDUCATIVA PRIVADA")
    run.font.name = 'Calibri'; run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*role_color)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("SANTÍSIMA CRUZ")
    run2.font.name = 'Calibri'; run2.font.size = Pt(18); run2.font.bold = True
    run2.font.color.rgb = RGBColor(*role_color)

    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sep.add_run("─" * 45)
    s_run.font.color.rgb = RGBColor(*role_color)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(20)
    run3 = p3.add_run(title)
    run3.font.name = 'Calibri'; run3.font.size = Pt(24); run3.font.bold = True
    run3.font.color.rgb = RGBColor(*role_color)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(subtitle)
    run4.font.name = 'Calibri'; run4.font.size = Pt(14); run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x1F, 0x7A, 0x8C)

    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sep2.add_run("─" * 45)
    s2.font.color.rgb = RGBColor(*role_color)

    doc.add_paragraph()
    info_items = [
        ("Versión:", version),
        ("Año lectivo:", "2026"),
        ("Fecha:", datetime.datetime.now().strftime("%d/%m/%Y")),
    ]
    for label, val in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + " "); r1.font.bold = True; r1.font.size = Pt(11)
        r2 = p.add_run(val); r2.font.size = Pt(11)
    doc.add_page_break()

# ════════════════════════════════════════════════════════════
#  CONTENIDO POR MÓDULO
# ════════════════════════════════════════════════════════════

def write_intro(doc):
    doc.add_heading("Introducción", level=1)
    body(doc, "Este manual ha sido elaborado para guiar a todos los usuarios del Sistema Educativo de la I.E.P. Santísima Cruz en el uso correcto de la plataforma digital. Encontrará instrucciones paso a paso para cada función disponible según su rol.")
    body(doc, "El sistema cuenta con tres perfiles de acceso:")
    bullet(doc, "Administrador (Director): acceso total para gestionar la institución.")
    bullet(doc, "Docente: acceso para revisar secciones, estudiantes y registrar notas.")
    bullet(doc, "Estudiante: acceso para consultar sus calificaciones y datos de matrícula.")
    tip_box(doc, "💡 RECOMENDACIÓN: Use Google Chrome o Microsoft Edge actualizados para la mejor experiencia. La plataforma está optimizada para navegadores modernos.", 'E8F4FD')
    doc.add_paragraph()

# ─── MÓDULO ADMINISTRADOR ───────────────────────────────────

def write_admin(doc):
    section_title_banner(doc, "PARTE 1 — MÓDULO DEL ADMINISTRADOR", '1F3564')
    doc.add_paragraph()

    # 1.1 INICIO DE SESIÓN
    doc.add_heading("1.1  Inicio de Sesión", level=2)
    body(doc, "El inicio de sesión es la puerta de entrada a la plataforma. Solo el administrador con credenciales válidas puede ingresar al panel de administración.")
    doc.add_paragraph()

    step(doc, 1, "Abra su navegador web (Google Chrome o Microsoft Edge) y diríjase a la dirección del sistema proporcionada por el responsable técnico (por ejemplo: https://sistema.colegio.pe).")
    screenshot_placeholder(doc, "Pantalla del navegador mostrando la URL del sistema antes de cargar la página de login")

    step(doc, 2, "Aparecerá la pantalla de inicio de sesión. Verá dos campos: \"DNI o Correo\" y \"Contraseña\".")
    screenshot_placeholder(doc, "Página de inicio de sesión completa, con los campos 'DNI o Correo' y 'Contraseña' visibles")

    step(doc, 3, "En el campo \"DNI o Correo\", ingrese su número de DNI o la dirección de correo electrónico registrada en el sistema.")

    step(doc, 4, "En el campo \"Contraseña\", ingrese su contraseña. Puede hacer clic en el ícono del ojo (👁) a la derecha del campo para mostrar u ocultar los caracteres.")
    screenshot_placeholder(doc, "Campo de contraseña con el ícono de mostrar/ocultar contraseña resaltado")

    step(doc, 5, "De forma opcional, marque la casilla \"Recordarme\" si desea que el sistema no le pida ingresar sus datos en el próximo acceso desde el mismo equipo.")

    step(doc, 6, "Haga clic en el botón dorado \"Ingresar\" para acceder.")
    screenshot_placeholder(doc, "Botón 'Ingresar' resaltado en la página de login")

    step(doc, 7, "Si sus datos son correctos, el sistema lo redirigirá automáticamente al Panel de Administración.")
    screenshot_placeholder(doc, "Panel de Administración (Admin Dashboard) recién cargado con las 3 tarjetas de resumen: Estudiantes Activos, Total de Docentes, Aulas Activas")

    tip_box(doc, "⚠️ Si ve el mensaje 'Estas credenciales no coinciden con nuestros registros', verifique que ingresó el DNI o correo correctamente. Si olvidó su contraseña, haga clic en '¿Olvidaste tu contraseña?' para restablecerla por correo electrónico.", 'FFF3E0')

    separator(doc)

    # 1.2 PANEL PRINCIPAL
    doc.add_heading("1.2  Panel Principal del Administrador (Dashboard)", level=2)
    body(doc, "Al ingresar, el administrador verá el panel principal con un resumen visual del estado actual de la institución.")
    body(doc, "El panel muestra:")
    bullet(doc, "Tarjeta Azul — «Estudiantes Activos»: muestra el número total de alumnos matriculados en el año lectivo vigente.")
    bullet(doc, "Tarjeta Verde — «Total de Docentes»: cantidad de docentes registrados en el sistema.")
    bullet(doc, "Tarjeta Morada — «Aulas Activas»: número de secciones configuradas para el año lectivo.")
    bullet(doc, "Panel «Recaudación (Mes)»: muestra el porcentaje de familias al día con los pagos del mes seleccionado. Puede cambiar el mes y año usando los selectores desplegables.")
    bullet(doc, "Panel «Consolidación de Notas»: barra de progreso que indica qué porcentaje de docentes ya registraron calificaciones en el bimestre en curso.")
    screenshot_placeholder(doc, "Dashboard completo del Administrador con todas las tarjetas de estadísticas y ambos paneles de Recaudación y Consolidación de Notas visibles")
    tip_box(doc, "💡 El dashboard se actualiza cada vez que recarga la página. Úselo al inicio del día para tener un panorama rápido del estado de la institución.", 'E8F4FD')

    separator(doc)

    # 1.3 GESTIÓN INSTITUCIONAL
    doc.add_heading("1.3  Gestión Institucional", level=2)
    body(doc, "Antes de iniciar el año lectivo, el administrador debe configurar la estructura académica de la institución. Este proceso se realiza una sola vez por año.")

    doc.add_heading("1.3.1  Configurar el Año Lectivo", level=3)
    body(doc, "El año lectivo es el período académico que agrupa todas las actividades: matrículas, bimestres y calificaciones.")

    step(doc, 1, "En el menú lateral o superior de navegación, busque y haga clic en \"Años Lectivos\".")
    screenshot_placeholder(doc, "Menú de navegación con la opción 'Años Lectivos' resaltada")

    step(doc, 2, "Verá la lista de años lectivos registrados. Si es el primer año, la lista estará vacía.")
    screenshot_placeholder(doc, "Pantalla de Años Lectivos con el listado y el formulario para crear un nuevo año")

    step(doc, 3, "Para crear un nuevo año, complete el campo «Año» con el año académico (ejemplo: 2026) y haga clic en «Crear Año Lectivo».")

    step(doc, 4, "Una vez creado, haga clic en el botón «Activar» junto al año que corresponda al período actual. Solo puede haber un año activo a la vez. El año activo aparecerá marcado en verde.")
    screenshot_placeholder(doc, "Lista de años lectivos con el botón 'Activar' resaltado y el año activo marcado con una etiqueta verde")
    tip_box(doc, "⚠️ IMPORTANTE: Si activa un año diferente, todas las operaciones del sistema (importaciones, notas, bimestres) pasarán a operar bajo ese año. Realice este cambio solo cuando inicie un nuevo período académico.", 'FFE8E8')

    doc.add_heading("1.3.2  Gestionar Grados y Secciones", level=3)
    body(doc, "Los grados y secciones definen la estructura de aulas del colegio para el año lectivo activo. Cada combinación de Grado + Sección + Año Lectivo forma un aula única.")

    step(doc, 1, "En el menú de navegación, haga clic en \"Grados y Secciones\".")
    screenshot_placeholder(doc, "Pantalla de Grados y Secciones con la lista completa de aulas del año activo, mostrando grado, sección, tutor asignado y número de estudiantes")

    step(doc, 2, "Para crear una nueva sección, haga clic en el botón «+ Nueva Sección» o «Agregar Grado-Sección».")

    step(doc, 3, "En el formulario que aparece, seleccione:")
    bullet(doc, "El Nivel (Primaria o Secundaria).", 1)
    bullet(doc, "El Grado (por ejemplo: 1er Grado, 2do Grado, etc.).", 1)
    bullet(doc, "La Sección (por ejemplo: A, B, C).", 1)
    screenshot_placeholder(doc, "Formulario de creación de nueva sección con los desplegables de Nivel, Grado y Sección")

    step(doc, 4, "Haga clic en «Guardar». La nueva sección aparecerá en la lista.")

    step(doc, 5, "Para asignar el docente Tutor y Co-Tutor a la sección, haga clic en el ícono de edición (✏️) o en el nombre de la sección, y luego seleccione al docente correspondiente en el campo «Tutor».")
    screenshot_placeholder(doc, "Formulario de edición de sección con los selectores de Tutor y Co-Tutor desplegados, mostrando la lista de docentes disponibles")
    tip_box(doc, "💡 El Tutor es el docente principal a cargo de la sección. Tiene acceso a todas las notas del aula y puede importar calificaciones. Asegúrese de asignarlo antes de que los docentes comiencen a registrar notas.", 'E8F4FD')

    separator(doc)

    # 1.4 GESTIÓN DE BIMESTRES
    doc.add_heading("1.4  Gestión de Bimestres", level=2)
    body(doc, "Los bimestres son los cuatro períodos de evaluación del año escolar. El administrador controla su apertura y cierre.")

    step(doc, 1, "En el menú de navegación, haga clic en \"Bimestres\".")
    screenshot_placeholder(doc, "Pantalla de Bimestres con la lista de los bimestres del año activo, mostrando número, fechas, estado (Abierto/Cerrado) y botones de acción")

    step(doc, 2, "Para crear el primer bimestre, haga clic en «+ Aperturar Bimestre». Complete los campos:")
    bullet(doc, "Fecha de Inicio: primer día del bimestre.", 1)
    bullet(doc, "Fecha de Fin: último día del bimestre.", 1)
    screenshot_placeholder(doc, "Formulario de creación de bimestre con los campos Fecha de Inicio y Fecha de Fin")

    step(doc, 3, "Haga clic en «Guardar». El bimestre se creará en estado «Abierto» y los docentes podrán registrar notas.")

    step(doc, 4, "Para cerrar un bimestre al finalizar el período, haga clic en «Cerrar Bimestre». El sistema le mostrará una pantalla de confirmación con un resumen de qué secciones tienen estudiantes sin notas registradas.")
    screenshot_placeholder(doc, "Pantalla de confirmación de cierre de bimestre, mostrando la tabla resumen con el nombre de cada sección, total de estudiantes y cantidad de estudiantes sin notas")

    step(doc, 5, "Revise el resumen. Si todo está correcto, haga clic en «Confirmar Cierre». El sistema calculará automáticamente los promedios bimestrales de todos los estudiantes.")
    tip_box(doc, "⚠️ Una vez cerrado un bimestre, los docentes NO podrán importar ni modificar notas de ese período. Verifique que todos los docentes hayan registrado las calificaciones antes de cerrar.", 'FFE8E8')

    step(doc, 6, "Si necesita reabrir un bimestre cerrado (por corrección de notas), haga clic en «Reabrir Bimestre» en la fila correspondiente.")
    screenshot_placeholder(doc, "Fila de bimestre cerrado con el botón 'Reabrir Bimestre' resaltado")

    step(doc, 7, "Para publicar las notas y que los estudiantes las puedan ver en su portal, haga clic en «Publicar Notas», seleccione el nivel (Primaria, Secundaria o Ambos) y confirme.")
    screenshot_placeholder(doc, "Modal o sección de publicación de notas con los botones de Primaria, Secundaria y Ambos")

    separator(doc)

    # 1.5 IMPORTACIÓN MASIVA
    doc.add_heading("1.5  Importación Masiva desde Excel (SIAGIE)", level=2)
    body(doc, "El sistema permite importar grandes cantidades de datos desde archivos Excel exportados del sistema SIAGIE del Ministerio de Educación, o desde cualquier planilla Excel con el formato adecuado. Esta función es la más usada al inicio del año escolar para registrar estudiantes, apoderados y notas.")

    doc.add_heading("1.5.1  Cómo Preparar el Archivo Excel", level=3)
    body(doc, "El archivo Excel que usted cargue puede tener cualquier nombre de columnas: el sistema le permitirá mapearlas manualmente. Sin embargo, se recomienda que el archivo contenga al menos:")
    bullet(doc, "Una columna con el DNI del estudiante.")
    bullet(doc, "Columnas con los nombres y apellidos (pueden estar separados o en una sola columna).")
    bullet(doc, "Fecha de nacimiento y sexo (opcional, pero recomendado).")
    tip_box(doc, "💡 El sistema soporta archivos .xlsx y .xls de Excel. El tamaño máximo permitido es de 20 MB. Los archivos exportados directamente desde SIAGIE son completamente compatibles.", 'E8F4FD')

    doc.add_heading("1.5.2  Proceso de Importación Paso a Paso", level=3)
    step(doc, 1, "En el menú de navegación, haga clic en «Importaciones» o «Importar desde Excel».")
    screenshot_placeholder(doc, "Pantalla de Importaciones con el historial de importaciones anteriores en la parte inferior y el formulario de nueva importación en la parte superior")

    step(doc, 2, "En el formulario de importación, seleccione el «Tipo de Importación» desde el menú desplegable:")
    bullet(doc, "Estudiantes: crea o actualiza los datos de alumnos y genera sus cuentas de acceso.", 1)
    bullet(doc, "Directorio (SIAGIE): importa la lista oficial de matrícula del SIAGIE con todos los datos del estudiante.", 1)
    bullet(doc, "Apoderados/Padres: importa los datos de los padres o tutores vinculados a los estudiantes.", 1)
    bullet(doc, "Notas Bimestrales: carga las calificaciones de un bimestre específico.", 1)
    screenshot_placeholder(doc, "Desplegable de Tipo de Importación con las opciones visibles: Estudiantes, Directorio SIAGIE, Apoderados, Notas Bimestrales")

    step(doc, 3, "Seleccione el Nivel Educativo (Primaria o Secundaria), el Grado y la Sección de destino para los datos que está importando.")

    step(doc, 4, "Haga clic en el área de carga de archivo (o «Seleccionar Archivo») y elija el archivo Excel desde su computadora.")
    screenshot_placeholder(doc, "Área de arrastre/subida de archivo con un archivo .xlsx ya seleccionado, mostrando el nombre del archivo")

    step(doc, 5, "Haga clic en «Vista Previa». El sistema leerá el archivo y mostrará una tabla de previsualización con las primeras filas del archivo.")
    screenshot_placeholder(doc, "Pantalla de vista previa con la tabla de datos del Excel mostrada y los selectores de mapeo de columnas en la fila de encabezado")

    step(doc, 6, "En la fila de encabezado de la tabla, aparecerán menús desplegables sobre cada columna. Usando estos menús, indique al sistema qué representa cada columna de su Excel:")
    bullet(doc, "Si la columna dice «DNI» → seleccione «DNI» en el desplegable.", 1)
    bullet(doc, "Si la columna dice «Nombres» → seleccione «Nombres».", 1)
    bullet(doc, "Si una columna no es necesaria → déjela en «— Ignorar columna —».", 1)
    screenshot_placeholder(doc, "Fila de mapeo de columnas con un desplegable abierto mostrando las opciones: DNI, Nombres, Apellido Paterno, Apellido Materno, Fecha de Nacimiento, Sexo, — Ignorar columna —")

    step(doc, 7, "Revise los datos en la tabla de previsualización. Si detecta filas con datos incorrectos o que no desea importar, desmarque la casilla de verificación (☑) al inicio de esa fila.")
    screenshot_placeholder(doc, "Tabla de vista previa con varias filas, resaltando la casilla de verificación en la primera columna y una fila desmarcada")

    step(doc, 8, "Una vez que el mapeo de columnas sea correcto y haya revisado los datos, haga clic en el botón «Confirmar Importación».")
    screenshot_placeholder(doc, "Botón 'Confirmar Importación' resaltado en la parte inferior de la pantalla de vista previa")

    step(doc, 9, "El sistema procesará el archivo. Al finalizar, mostrará un mensaje de resultado con:")
    bullet(doc, "Número de registros importados exitosamente (en verde).", 1)
    bullet(doc, "Número de registros con error y descripción de cada fila problemática (en rojo).", 1)
    screenshot_placeholder(doc, "Pantalla de resultado de importación mostrando '45 estudiantes importados correctamente' en verde y una lista de errores por fila en rojo si los hubiera")

    step(doc, 10, "Si necesita deshacer una importación completa (por ejemplo, si cargó el archivo equivocado), vaya al historial de importaciones en la parte inferior de la pantalla. Localice la importación y haga clic en «Revertir».")
    screenshot_placeholder(doc, "Historial de importaciones con la tabla que muestra fecha, tipo, archivo, cantidad importada y botón 'Revertir' resaltado en la última importación")
    tip_box(doc, "⚠️ La función 'Revertir' elimina TODOS los registros creados durante esa importación. Úsela con cuidado, ya que no se puede deshacer.", 'FFE8E8')

    separator(doc)

    # 1.6 GESTIÓN DE DOCENTES
    doc.add_heading("1.6  Gestión de Docentes", level=2)
    body(doc, "El administrador puede registrar nuevos docentes, ver su perfil completo y gestionar sus asignaciones a cursos y secciones.")

    doc.add_heading("1.6.1  Registrar un Nuevo Docente", level=3)
    step(doc, 1, "En el menú de navegación, haga clic en «Docentes».")
    screenshot_placeholder(doc, "Pantalla de lista de docentes con la tabla mostrando nombre, DNI, tipo (Polidocente/Especialista), nivel y botones de acción")

    step(doc, 2, "Haga clic en «+ Nuevo Docente».")
    step(doc, 3, "Complete el formulario con los datos del docente:")
    bullet(doc, "DNI (obligatorio): se usará como contraseña inicial.", 1)
    bullet(doc, "Nombres, Apellido Paterno y Apellido Materno.", 1)
    bullet(doc, "Celular (opcional).", 1)
    bullet(doc, "Nivel Educativo: Primaria o Secundaria.", 1)
    bullet(doc, "Tipo de docente:", 1)
    bullet(doc, "Polidocente: enseña todas las materias a un mismo grado (típico en 1ro a 4to de Primaria).", 2)
    bullet(doc, "Especialista: enseña uno o más cursos específicos en diferentes secciones.", 2)
    screenshot_placeholder(doc, "Formulario de creación de nuevo docente con todos los campos completos y el selector de Tipo (Polidocente/Especialista) desplegado")

    step(doc, 4, "Haga clic en «Guardar». El sistema creará automáticamente la cuenta de acceso del docente con su DNI como usuario y contraseña inicial.")
    tip_box(doc, "💡 El docente recibirá el mensaje de que debe cambiar su contraseña en el primer ingreso al sistema. Infórmele que su credencial inicial es su número de DNI.", 'E8F4FD')

    doc.add_heading("1.6.2  Asignar Docentes a Cursos y Secciones", level=3)
    body(doc, "Después de registrar al docente, debe asignarle las secciones y cursos que dictará durante el año lectivo.")

    step(doc, 1, "En la lista de docentes, haga clic en el nombre del docente o en el botón «Ver Perfil».")
    screenshot_placeholder(doc, "Perfil del docente con sus datos personales y la sección de 'Asignaciones del año lectivo' en la parte inferior")

    step(doc, 2, "En el perfil del docente, ubique la sección «Asignar a Sección/Curso» o el botón «+ Nueva Asignación».")

    step(doc, 3, "Para un docente Especialista, seleccione:")
    bullet(doc, "La Sección (Grado + Sección, ej: «1er Grado - Sección A»).", 1)
    bullet(doc, "El Curso que dictará en esa sección (ej: Matemática, Comunicación).", 1)
    screenshot_placeholder(doc, "Formulario de asignación de docente especialista con el selector de Sección y el selector de Curso desplegado mostrando la lista de cursos disponibles")

    step(doc, 4, "Para un docente Polidocente, seleccione únicamente la Sección. El sistema entenderá que dictará todos los cursos de ese grado.")
    screenshot_placeholder(doc, "Formulario de asignación de docente polidocente, mostrando solo el selector de Sección (sin selector de curso)")

    step(doc, 5, "Haga clic en «Asignar». Si el curso ya tiene otro docente asignado en esa sección, el sistema mostrará un mensaje de advertencia indicando el conflicto y no permitirá la asignación duplicada.")
    screenshot_placeholder(doc, "Mensaje de error/advertencia que dice 'El docente [Nombre] ya tiene ese curso en esa sección' cuando hay un conflicto de asignación")
    tip_box(doc, "💡 Un mismo docente puede tener múltiples asignaciones: puede dictar Matemática en 3ro A y Matemática en 3ro B simultáneamente. Repita el proceso para cada combinación.", 'E8F4FD')

    separator(doc)

    # 1.7 GESTIÓN DE ESTUDIANTES
    doc.add_heading("1.7  Gestión de Estudiantes", level=2)
    body(doc, "Además de la importación masiva, el administrador puede gestionar estudiantes individualmente.")

    step(doc, 1, "En el menú de navegación, haga clic en «Estudiantes».")
    screenshot_placeholder(doc, "Pantalla de lista de estudiantes con buscador, filtros por grado/sección y tabla con columnas: Nombre, DNI, Grado, Sección, Estado y Acciones")

    step(doc, 2, "Use la barra de búsqueda para encontrar un estudiante por nombre o DNI.")
    step(doc, 3, "Haga clic en el nombre del estudiante para ver su perfil completo, incluyendo datos personales, información de su apoderado y su historial de matrículas.")
    screenshot_placeholder(doc, "Perfil detallado del estudiante mostrando sus datos personales, datos del apoderado y estado de matrícula actual")

    step(doc, 4, "Si necesita mover a un estudiante a otra sección (por ejemplo, cambio de sección), busque el botón «Cambiar de Sección» y seleccione el nuevo grado y sección.")
    tip_box(doc, "⚠️ Al mover a un estudiante de sección, sus notas ya registradas permanecen intactas. Solo se actualiza su matrícula activa.", 'FFF3E0')

    separator(doc)

    # 1.8 MENSUALIDADES
    doc.add_heading("1.8  Control de Mensualidades", level=2)
    body(doc, "El módulo de mensualidades permite llevar el control de los pagos mensuales de cada familia.")

    step(doc, 1, "En el menú de navegación, haga clic en «Mensualidades».")
    screenshot_placeholder(doc, "Pantalla de mensualidades con filtros de mes, año y sección, y la tabla de estudiantes con su estado de pago: Al día, Pendiente, o Pagado")

    step(doc, 2, "Use los filtros de Mes, Año y Sección para visualizar el estado de pagos del período que necesita.")
    step(doc, 3, "Para registrar un pago, haga clic en el botón «Registrar Pago» junto al nombre del estudiante.")
    screenshot_placeholder(doc, "Fila de estudiante en la tabla de mensualidades con el botón 'Registrar Pago' resaltado y el estado actual 'Pendiente'")
    step(doc, 4, "El sistema actualizará el estado del estudiante a «Pagado» y registrará la fecha del pago.")

    separator(doc)

    # 1.9 REPORTES
    doc.add_heading("1.9  Reportes y Rendimiento Académico", level=2)
    body(doc, "El módulo de Rendimiento le permite visualizar y exportar informes del desempeño académico de los estudiantes.")

    step(doc, 1, "En el menú de navegación, haga clic en «Rendimiento».")
    screenshot_placeholder(doc, "Pantalla principal de Rendimiento con las pestañas disponibles: Por Sección, Por Docente, Estudiantes Críticos y Exportar")

    step(doc, 2, "Seleccione la pestaña del tipo de reporte que necesita:")
    bullet(doc, "«Por Sección»: muestra el promedio general de cada sección por bimestre.", 1)
    bullet(doc, "«Por Docente»: muestra el avance de calificaciones registradas por cada docente.", 1)
    bullet(doc, "«Estudiantes Críticos»: lista a los estudiantes con calificación C (En Inicio) en uno o más cursos.", 1)
    screenshot_placeholder(doc, "Pestaña 'Estudiantes Críticos' activa, mostrando la lista de alumnos con C en algún curso, con columnas: Nombre, Grado, Sección y Cursos con C")

    step(doc, 3, "Para exportar el reporte a Excel, haga clic en el botón «Exportar a Excel». Se descargará automáticamente un archivo .xlsx en su computadora.")
    screenshot_placeholder(doc, "Botón 'Exportar a Excel' resaltado en la esquina de la pantalla de reportes")

# ─── MÓDULO DOCENTE ─────────────────────────────────────────

def write_docente(doc):
    section_title_banner(doc, "PARTE 2 — MÓDULO DEL DOCENTE", '0F5132')
    doc.add_paragraph()

    # 2.1 PRIMER INGRESO
    doc.add_heading("2.1  Primer Ingreso — Cambio de Contraseña Obligatorio", level=2)
    body(doc, "Al iniciar sesión por primera vez, el sistema detectará que su cuenta tiene la contraseña predeterminada (su número de DNI) y le pedirá obligatoriamente que la cambie por una más segura. No podrá acceder a ninguna función del sistema sin completar este paso.")

    step(doc, 1, "Abra su navegador y vaya a la dirección del sistema (proporcionada por el director o administrador).")
    step(doc, 2, "En la pantalla de inicio de sesión, ingrese su DNI en el campo «DNI o Correo» y su número de DNI también en el campo «Contraseña» (es la contraseña inicial).")
    screenshot_placeholder(doc, "Pantalla de inicio de sesión con el campo DNI completado y el campo Contraseña completado con el DNI como contraseña inicial")

    step(doc, 3, "Haga clic en «Ingresar». El sistema lo redirigirá automáticamente a la pantalla de cambio obligatorio de contraseña.")
    screenshot_placeholder(doc, "Pantalla de cambio obligatorio de contraseña con el cuadro azul de instrucciones que dice 'Por razones de seguridad, debes cambiar tu contraseña antes de continuar' y los requisitos listados")

    step(doc, 4, "Lea los requisitos de la nueva contraseña en el cuadro azul informativo:")
    bullet(doc, "Mínimo 8 caracteres.", 1)
    bullet(doc, "Al menos una letra mayúscula (A-Z) y una minúscula (a-z).", 1)
    bullet(doc, "Al menos un número (0-9).", 1)
    bullet(doc, "Al menos un carácter especial como: @, $, !, %, *, ?, &.", 1)

    step(doc, 5, "Escriba su nueva contraseña en el campo «Nueva Contraseña». Elija una contraseña que pueda recordar fácilmente pero que sea difícil de adivinar.")
    step(doc, 6, "Repita la misma contraseña en el campo «Confirmar Nueva Contraseña».")
    screenshot_placeholder(doc, "Formulario de cambio de contraseña con ambos campos (Nueva Contraseña y Confirmar Nueva Contraseña) completados, con los puntos de contraseña visibles")

    step(doc, 7, "Haga clic en el botón amarillo «Cambiar Contraseña».")
    step(doc, 8, "El sistema aceptará su nueva contraseña y lo llevará directamente a su Panel Docente.")
    screenshot_placeholder(doc, "Panel Principal del Docente (Mi Panel Docente) recién cargado, mostrando la tarjeta verde con el nombre completo del docente, su DNI, nivel y tipo")
    tip_box(doc, "💡 Si no está listo para cambiar la contraseña en ese momento, puede hacer clic en el enlace 'Omitir por ahora' en la parte inferior del formulario. Sin embargo, el sistema le solicitará el cambio nuevamente en el próximo inicio de sesión.", 'E8F4FD')

    separator(doc)

    # 2.2 PANEL PRINCIPAL DOCENTE
    doc.add_heading("2.2  Panel Principal del Docente", level=2)
    body(doc, "Después de iniciar sesión, verá su panel personal con toda su información institucional.")
    body(doc, "La tarjeta principal de su perfil muestra:")
    bullet(doc, "Su nombre completo (en formato: Apellido Paterno, Apellido Materno, Nombres).")
    bullet(doc, "Un badge que indica si es «Polidocente» o «Especialista».")
    bullet(doc, "Su DNI, Nivel educativo (Primaria/Secundaria), celular de contacto y tipo de docente.")
    bullet(doc, "Si es especialista: una sección adicional con los cursos de su especialidad.")
    screenshot_placeholder(doc, "Panel docente completo: tarjeta de perfil con encabezado verde, nombre, badge de tipo, y grilla con DNI, Nivel, Celular y Tipo")
    tip_box(doc, "⚠️ Si su panel muestra el mensaje 'Perfil no encontrado', significa que el administrador aún no ha vinculado su cuenta con un perfil de docente. Comuníquese con el director o el responsable del sistema.", 'FFF3E0')

    separator(doc)

    # 2.3 MIS SECCIONES
    doc.add_heading("2.3  Ver Mis Secciones Asignadas", level=2)
    body(doc, "Para ver las secciones donde dicta clases, siga estos pasos:")

    step(doc, 1, "En el menú de navegación, haga clic en «Mis Secciones» o «Secciones a Cargo».")
    screenshot_placeholder(doc, "Pantalla 'Secciones a Cargo' con tarjetas de cada sección asignada: cada tarjeta muestra el nombre del grado, la sección (ej: A, B) y el nivel. Las secciones donde es tutor tienen una banda dorada superior que dice 'Sección a cargo (Tutor)'")

    step(doc, 2, "Verá una cuadrícula con tarjetas, una por cada sección asignada. Las tarjetas con una franja dorada en la parte superior indican que usted es el «Tutor» de esa sección.")

    step(doc, 3, "Haga clic sobre cualquier tarjeta de sección para ver la lista de estudiantes de esa sección.")
    screenshot_placeholder(doc, "Pantalla de lista de estudiantes de una sección específica, con columnas: Nombre completo, DNI, y botón 'Ver Notas'")
    tip_box(doc, "💡 Como Tutor de una sección, usted puede ver las notas de TODOS los cursos de sus estudiantes, no solo los cursos que dicta. Como especialista sin tutoría, solo verá las notas de su curso.", 'E8F4FD')

    separator(doc)

    # 2.4 VER NOTAS DE UN ESTUDIANTE
    doc.add_heading("2.4  Ver y Consultar Notas de un Estudiante", level=2)
    body(doc, "Desde la lista de estudiantes de su sección, puede consultar el historial completo de calificaciones.")

    step(doc, 1, "En la lista de estudiantes, haga clic en el botón «Ver Notas» junto al nombre del estudiante que desea consultar.")
    screenshot_placeholder(doc, "Lista de estudiantes de la sección con el botón 'Ver Notas' resaltado junto a un nombre de estudiante")

    step(doc, 2, "Verá la «Libreta Digital» del estudiante: una ficha con su nombre, DNI, grado, sección y año lectivo en la parte superior.")
    screenshot_placeholder(doc, "Ficha de identificación del estudiante en la parte superior de la pantalla de notas, con su avatar con iniciales, nombre completo, DNI, grado-sección y año lectivo")

    step(doc, 3, "Debajo de la ficha, verá una tabla por cada curso. Cada tabla tiene:")
    bullet(doc, "La columna «Competencias»: lista las competencias o capacidades evaluadas en ese curso.", 1)
    bullet(doc, "Las columnas «B1», «B2», «B3», «B4»: muestran la nota de cada bimestre para esa competencia.", 1)
    screenshot_placeholder(doc, "Tabla de notas de un curso (ej: Matemática) con filas de competencias y columnas B1, B2, B3, B4. Las notas se muestran con colores: AD en morado/azul, A en verde, B en amarillo, C en rojo")

    step(doc, 4, "Las notas se muestran con un código de colores para identificar rápidamente el nivel de logro:")
    bullet(doc, "AD (morado/azul) — Logro Destacado: el estudiante supera los objetivos.", 1)
    bullet(doc, "A (verde) — Logro Esperado: el estudiante cumple los objetivos.", 1)
    bullet(doc, "B (amarillo/ámbar) — En Proceso: el estudiante está avanzando hacia los objetivos.", 1)
    bullet(doc, "C (rojo) — En Inicio: el estudiante requiere apoyo para alcanzar los objetivos.", 1)
    bullet(doc, "— (gris) — Sin nota: aún no se ha registrado calificación para ese período.", 1)
    screenshot_placeholder(doc, "Leyenda de colores de notas o detalle de una tabla de notas mostrando los 4 colores diferentes: AD en azul, A en verde, B en ámbar, C en rojo")

    step(doc, 5, "Use el botón «← Atrás» en la parte superior izquierda para volver a la lista de estudiantes de la sección.")

    separator(doc)

    # 2.5 IMPORTAR NOTAS
    doc.add_heading("2.5  Importar Notas Bimestrales (Docente Tutor)", level=2)
    body(doc, "Los docentes que son Tutores de una sección pueden importar las notas bimestrales directamente desde un archivo Excel. Esta función solo está disponible para tutores.")

    step(doc, 1, "En el menú de navegación, haga clic en «Importar Notas».")
    screenshot_placeholder(doc, "Pantalla de importación de notas del docente con el formulario de subida de archivo y el selector de bimestre")

    step(doc, 2, "Seleccione el Bimestre para el cual está importando las notas (B1, B2, B3 o B4). Solo podrá importar a bimestres con estado «Abierto».")
    screenshot_placeholder(doc, "Selector de Bimestre con las opciones B1, B2, B3, B4 y el estado de cada uno indicado")

    step(doc, 3, "Seleccione o arrastre el archivo Excel con las notas al área de carga.")
    step(doc, 4, "Haga clic en «Vista Previa» para verificar que el archivo está siendo leído correctamente.")
    screenshot_placeholder(doc, "Vista previa del archivo de notas mostrando columnas: Estudiante, Competencia, y la nota (AD/A/B/C)")

    step(doc, 5, "Revise los datos. Si todo es correcto, haga clic en «Confirmar Importación».")
    step(doc, 6, "El sistema procesará el archivo y mostrará el resumen de notas importadas.")
    tip_box(doc, "⚠️ Si el bimestre está CERRADO, no podrá importar notas. Comuníquese con el administrador para que lo reabra si necesita hacer correcciones.", 'FFE8E8')

# ─── MÓDULO ESTUDIANTE ──────────────────────────────────────

def write_estudiante(doc):
    section_title_banner(doc, "PARTE 3 — MÓDULO DEL ESTUDIANTE", '3730A3')
    doc.add_paragraph()

    # 3.1 PRIMER INGRESO
    doc.add_heading("3.1  Primer Ingreso — Cambio de Contraseña Obligatorio", level=2)
    body(doc, "Al igual que los docentes, el primer acceso de cada estudiante requiere el cambio obligatorio de la contraseña predeterminada.")

    step(doc, 1, "Abra el navegador y acceda a la dirección del sistema proporcionada por el colegio.")
    step(doc, 2, "En la pantalla de inicio de sesión, ingrese su número de DNI en el campo «DNI o Correo».")
    step(doc, 3, "En el campo «Contraseña», ingrese nuevamente su número de DNI. Esta es su contraseña inicial asignada por el colegio.")
    screenshot_placeholder(doc, "Pantalla de login con el campo DNI completado con el número de documento del estudiante")

    step(doc, 4, "Haga clic en «Ingresar». El sistema lo llevará a la pantalla de cambio obligatorio de contraseña.")
    screenshot_placeholder(doc, "Pantalla de cambio obligatorio de contraseña con el mensaje 'Por razones de seguridad, debes cambiar tu contraseña' y los requisitos listados")

    step(doc, 5, "Cree su nueva contraseña. Debe cumplir los siguientes requisitos:")
    bullet(doc, "Mínimo 8 caracteres.", 1)
    bullet(doc, "Al menos una mayúscula, una minúscula, un número y un carácter especial (@, $, !, %).", 1)

    step(doc, 6, "Escriba la nueva contraseña en «Nueva Contraseña» y repítala en «Confirmar Nueva Contraseña».")
    step(doc, 7, "Haga clic en «Cambiar Contraseña». El sistema lo llevará a su Panel Escolar personal.")
    screenshot_placeholder(doc, "Panel Escolar del Estudiante ('Mi Panel Escolar') con la tarjeta de perfil azul-índigo mostrando el nombre del estudiante, su grado y sección en el encabezado")
    tip_box(doc, "⚠️ Guarde su contraseña en un lugar seguro. Si la olvida, deberá pedirle al administrador del colegio que se la restablezca.", 'FFE8E8')

    separator(doc)

    # 3.2 PANEL DEL ESTUDIANTE
    doc.add_heading("3.2  Panel Principal del Estudiante (Mi Panel Escolar)", level=2)
    body(doc, "Al ingresar al sistema, verá su panel personal con todos sus datos de matrícula.")
    body(doc, "La tarjeta de perfil en la parte superior muestra:")
    bullet(doc, "Su nombre completo (en formato Apellido, Nombres).")
    bullet(doc, "Un badge con su Grado y Sección actual (ej: «3er Grado — Sección B»).")
    screenshot_placeholder(doc, "Tarjeta de encabezado del perfil del estudiante con gradiente azul-índigo, nombre completo en blanco y el badge con el grado y sección")

    body(doc, "Debajo del encabezado, verá una grilla con sus datos personales:")
    bullet(doc, "DNI.")
    bullet(doc, "Sexo (Masculino / Femenino).")
    bullet(doc, "Fecha de Nacimiento.")
    bullet(doc, "Nivel (Primaria / Secundaria).")
    bullet(doc, "Estado (Activo / Retirado).")
    bullet(doc, "Año Lectivo.")
    bullet(doc, "Grado y Sección.")
    screenshot_placeholder(doc, "Grilla de datos personales del estudiante con 8 campos: DNI, Sexo, Fecha de Nacimiento, Nivel, Estado, Año Lectivo, Grado y Sección")

    body(doc, "Si el colegio registró los datos de su apoderado, verá una sección adicional en la parte inferior con la información del padre, madre o tutor legal:")
    bullet(doc, "Nombre completo del apoderado.")
    bullet(doc, "DNI del apoderado.")
    bullet(doc, "Parentesco (Padre, Madre, Tutor).")
    bullet(doc, "Número de teléfono.")
    screenshot_placeholder(doc, "Sección 'Apoderado' en la parte inferior de la ficha del estudiante con los datos del padre/madre: nombre, DNI, parentesco y teléfono")
    tip_box(doc, "💡 Si sus datos personales o los de su apoderado son incorrectos, comuníquese con el área de administración del colegio para que los corrijan.", 'E8F4FD')

    separator(doc)

    # 3.3 VER NOTAS
    doc.add_heading("3.3  Consultar Mis Notas y Calificaciones", level=2)
    body(doc, "Cuando el administrador publique las notas de un bimestre, podrá consultarlas desde su portal.")

    step(doc, 1, "En el menú de navegación, busque y haga clic en «Mis Notas» o «Mi Libreta».")
    screenshot_placeholder(doc, "Menú de navegación del estudiante con la opción 'Mis Notas' resaltada")

    step(doc, 2, "Aparecerá su libreta digital con todas las notas registradas. Verá una tabla por cada curso, organizada con:")
    bullet(doc, "La primera columna «Competencias»: describe qué habilidad se evaluó.", 1)
    bullet(doc, "Las columnas B1, B2, B3, B4: muestran su calificación por bimestre.", 1)
    screenshot_placeholder(doc, "Libreta del estudiante con varias tablas de cursos (Matemática, Comunicación, etc.), cada una con sus competencias y las notas B1, B2, B3, B4 con colores")

    step(doc, 3, "Interprete sus calificaciones con el siguiente código:")
    bullet(doc, "AD — Logro Destacado: ¡Excelente! Superaste los objetivos del bimestre.", 1)
    bullet(doc, "A — Logro Esperado: Muy bien. Alcanzaste todos los objetivos.", 1)
    bullet(doc, "B — En Proceso: Vas por buen camino, pero debes reforzar algunos temas.", 1)
    bullet(doc, "C — En Inicio: Necesitas apoyo adicional. Habla con tu profesor o tutor.", 1)
    bullet(doc, "— (guión): Todavía no hay nota registrada para ese bimestre.", 1)
    screenshot_placeholder(doc, "Detalle de una tabla de notas del estudiante con diferentes calificaciones: AD en morado, A en verde, B en ámbar, C en rojo y guión en gris, para ilustrar la leyenda")

    tip_box(doc, "💡 Si no ves ninguna nota, es posible que el administrador aún no haya publicado los resultados del bimestre. Las notas solo son visibles para los estudiantes después de que el administrador active la publicación.", 'E8F4FD')

    separator(doc)

    # 3.4 PERFIL
    doc.add_heading("3.4  Mi Perfil y Datos de Cuenta", level=2)
    body(doc, "Puede acceder y editar algunos datos de su cuenta personal.")

    step(doc, 1, "Haga clic en su nombre o en el menú de usuario en la esquina superior derecha de la pantalla.")
    screenshot_placeholder(doc, "Esquina superior derecha del sistema con el menú de usuario desplegado, mostrando opciones: Mi Perfil, Cambiar Contraseña, Cerrar Sesión")

    step(doc, 2, "Seleccione «Perfil» o «Mi Cuenta». Podrá ver y actualizar su nombre de visualización y correo electrónico si lo tiene registrado.")

    step(doc, 3, "Para cambiar su contraseña de forma voluntaria (no solo cuando se le obliga), busque la opción «Cambiar Contraseña» dentro del perfil.")
    screenshot_placeholder(doc, "Sección de cambio de contraseña dentro del perfil del estudiante con los campos: Contraseña actual, Nueva contraseña y Confirmar contraseña")

    step(doc, 4, "Para cerrar sesión de forma segura, haga clic en «Cerrar Sesión» en el menú de usuario.")
    tip_box(doc, "🔒 SEGURIDAD: Siempre cierre sesión cuando use computadoras compartidas, como las del laboratorio del colegio. De esta forma, nadie más podrá ver sus datos.", 'FFE8E8')

# ────────────────────────────────────────────────────────────
# APÉNDICES COMUNES
# ────────────────────────────────────────────────────────────

def write_apendice(doc):
    doc.add_page_break()
    doc.add_heading("Apéndice A — Preguntas Frecuentes (FAQ)", level=1)

    faqs = [
        ("¿Qué hago si olvidé mi contraseña?",
         "En la pantalla de inicio de sesión, haga clic en '¿Olvidaste tu contraseña?'. "
         "Si tiene un correo electrónico registrado, recibirá un enlace para restablecerla. "
         "Si no tiene correo registrado, comuníquese con el administrador del colegio para que restablezca su acceso."),
        ("¿Por qué no puedo ver mis notas?",
         "Las notas solo son visibles para los estudiantes después de que el administrador "
         "publique los resultados del bimestre. Si el bimestre aún está en curso o el "
         "administrador no ha activado la publicación, las notas no estarán disponibles."),
        ("¿Puedo cambiar mi contraseña en cualquier momento?",
         "Sí. Vaya a su Perfil (clic en su nombre en la esquina superior derecha) y busque "
         "la opción 'Cambiar Contraseña'. Necesitará ingresar su contraseña actual para confirmar el cambio."),
        ("Como docente, ¿por qué no puedo ver las notas de ciertos cursos?",
         "Solo puede ver las notas de los cursos que tiene asignados. Si es tutor de una sección, "
         "puede ver todas las notas de esa sección. Si es especialista, solo verá las notas "
         "del curso que dicta en cada sección asignada."),
        ("¿Qué significa 'Bimestre Cerrado' y por qué no puedo importar notas?",
         "Un bimestre cerrado significa que el período de evaluación finalizó y el administrador "
         "calculó los promedios finales. Los docentes no pueden modificar notas de bimestres cerrados. "
         "Si necesita hacer una corrección, el administrador puede reabrir el bimestre temporalmente."),
        ("¿Puedo acceder al sistema desde mi celular?",
         "Sí. El sistema funciona en navegadores móviles (Chrome, Safari). Se recomienda usar "
         "el modo horizontal para ver mejor las tablas de notas."),
    ]

    for pregunta, respuesta in faqs:
        p = doc.add_paragraph()
        run_q = p.add_run(f"❓ {pregunta}")
        run_q.font.bold = True
        run_q.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
        body(doc, respuesta)
        doc.add_paragraph()

    doc.add_heading("Apéndice B — Escala de Calificaciones", level=1)
    body(doc, "El sistema utiliza la escala de calificación literal del Currículo Nacional de Educación Básica (CNEB) del Ministerio de Educación del Perú:")

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Calificación", "Descripción", "Significado"]
    header_colors = ['1F3564', '1F3564', '1F3564']
    row0 = table.rows[0]
    for i, (h, c) in enumerate(zip(headers, header_colors)):
        cell = row0.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), c)
        cell._tc.get_or_add_tcPr().append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        ("AD", "Logro Destacado", "El estudiante evidencia un nivel superior al esperado. Demuestra aprendizajes que van más allá del nivel de logro esperado."),
        ("A", "Logro Esperado", "El estudiante evidencia el nivel esperado respecto a la competencia, demostrando manejo satisfactorio en las tareas propuestas."),
        ("B", "En Proceso", "El estudiante está próximo o cerca al nivel esperado, para lo cual requiere acompañamiento durante un tiempo razonable."),
        ("C", "En Inicio", "El estudiante muestra progreso mínimo en una competencia. Evidencia dificultades y necesita mayor tiempo de acompañamiento e intervención."),
    ]
    row_fills = ['E8F0FF', 'E8FFE8', 'FFFCE8', 'FFE8E8']
    for i, (cal, desc, sig) in enumerate(rows_data):
        row = table.rows[i+1]
        row.cells[0].text = cal
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = desc
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        row.cells[2].text = sig
        for cell in row.cells:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), row_fills[i])
            cell._tc.get_or_add_tcPr().append(shd)

    doc.add_paragraph()
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_final.paragraph_format.space_before = Pt(30)
    run_final = p_final.add_run("— Fin del Manual —")
    run_final.font.name = 'Calibri'; run_final.font.size = Pt(12)
    run_final.font.italic = True; run_final.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ver = p_ver.add_run(f"Versión 1.0 | I.E.P. Santísima Cruz | {datetime.datetime.now().strftime('%Y')}")
    run_ver.font.name = 'Calibri'; run_ver.font.size = Pt(9)
    run_ver.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ════════════════════════════════════════════════════════════
#  GENERACIÓN DE LOS 4 ARCHIVOS
# ════════════════════════════════════════════════════════════

import os
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── 1. MANUAL GENERAL ─────────────────────────────────────────
print("Generando Manual_General_Directorio.docx ...")
doc_general = Document()
cover_page(doc_general, "MANUAL DE USUARIO GENERAL",
           "Guía Completa del Sistema Educativo — Todos los Roles",
           (0x1F, 0x35, 0x64))
setup_document(doc_general, (0x1F, 0x35, 0x64))
write_intro(doc_general)
doc_general.add_page_break()
write_admin(doc_general)
doc_general.add_page_break()
write_docente(doc_general)
doc_general.add_page_break()
write_estudiante(doc_general)
write_apendice(doc_general)
doc_general.save(os.path.join(OUTPUT_DIR, "Manual_General_Directorio.docx"))
print("  OK -> Manual_General_Directorio.docx")

# ── 2. MANUAL ADMINISTRADOR ───────────────────────────────────
print("Generando Manual_Administrador.docx ...")
doc_admin = Document()
cover_page(doc_admin, "MANUAL DEL ADMINISTRADOR",
           "Guía de Uso — Director / Personal Administrativo",
           (0x1F, 0x35, 0x64))
setup_document(doc_admin, (0x1F, 0x35, 0x64))
write_intro(doc_admin)
doc_admin.add_page_break()
write_admin(doc_admin)
write_apendice(doc_admin)
doc_admin.save(os.path.join(OUTPUT_DIR, "Manual_Administrador.docx"))
print("  OK -> Manual_Administrador.docx")

# ── 3. MANUAL DOCENTE ─────────────────────────────────────────
print("Generando Manual_Docente.docx ...")
doc_docente = Document()
cover_page(doc_docente, "MANUAL DEL DOCENTE",
           "Guía de Uso — Profesores y Tutores",
           (0x0F, 0x51, 0x32))
setup_document(doc_docente, (0x0F, 0x51, 0x32))
write_intro(doc_docente)
doc_docente.add_page_break()
write_docente(doc_docente)
write_apendice(doc_docente)
doc_docente.save(os.path.join(OUTPUT_DIR, "Manual_Docente.docx"))
print("  OK -> Manual_Docente.docx")

# ── 4. MANUAL ESTUDIANTE ──────────────────────────────────────
print("Generando Manual_Estudiante.docx ...")
doc_est = Document()
cover_page(doc_est, "MANUAL DEL ESTUDIANTE",
           "Guía de Uso — Alumnos de la I.E.P. Santísima Cruz",
           (0x37, 0x30, 0xA3))
setup_document(doc_est, (0x37, 0x30, 0xA3))
write_intro(doc_est)
doc_est.add_page_break()
write_estudiante(doc_est)
write_apendice(doc_est)
doc_est.save(os.path.join(OUTPUT_DIR, "Manual_Estudiante.docx"))
print("  OK -> Manual_Estudiante.docx")

print("\nTodos los manuales generados exitosamente en:", OUTPUT_DIR)
